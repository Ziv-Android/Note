#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import sys
import time
import os

from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import Cm
from docx.oxml.ns import qn
from docxcompose.composer import Composer

from PIL import Image
from pydocx import PyDocX
import collections


class WordReport:
    def __init__(self):
        self.sn = ""
        self.base_info = None
        self.start_time = 0
        self.report_file_path = None
        self.document = None
        self.compose_files = []

    def reset_flag(self):
        self.sn = ""
        self.start_time = 0
        self.report_file_path = None
        self.document = None

    def save_doc(self, filename=None):
        _doc_path = os.path.join(os.getcwd(), "doc")
        if not os.path.exists(_doc_path):
            os.makedirs(_doc_path)
        if filename is None:
            filename = ""
        file_create_time = time.strftime('%Y-%m-%d-%H-%M-%S', time.localtime((time.time())))
        self.report_file_path = os.path.join(_doc_path, f"测试报告-{filename}-{file_create_time}.docx")
        self.document.save(self.report_file_path)

    def add_compose_file(self, f=None):
        if f is not None:
            self.compose_files.append(f)
            return True
        if self.report_file_path is not None:
            self.compose_files.append(self.report_file_path)
            return True
        return False

    def get_report_file_path(self):
        return self.report_file_path

    def get_compose_files(self):
        return self.compose_files

    def clear_compose_file(self):
        self.compose_files.clear()

    def compose_doc(self, files, filename=None):
        if filename is None:
            filename = ""
        path = os.path.join(os.getcwd()) + "\\libutils\\default.docx"
        new_document = Document(path)
        composer = Composer(new_document)
        for fn in files:
            composer.append(Document(fn))
        _doc_path = os.path.join(os.getcwd(), "doc")
        if not os.path.exists(_doc_path):
            os.makedirs(_doc_path)

        file_create_time = time.strftime('%Y-%m-%d-%H-%M-%S', time.localtime((time.time())))
        self.report_file_path = os.path.join(_doc_path, f"测试报告-{filename}-{file_create_time}-compose.docx")
        composer.save(self.report_file_path)
        return self.report_file_path

    def doc2html(self, file_path=None):
        if file_path is not None:
            self.report_file_path = file_path
        if self.report_file_path is None:
            return ""
        _doc_path = os.path.join(os.getcwd(), "doc")
        _report_html_path = os.path.join(_doc_path, "index.html")
        with open(_report_html_path, "w", encoding="utf-8") as wf:
            html_content = PyDocX.to_html(self.report_file_path)
            wf.write(html_content)
        self.reset_flag()
        return _report_html_path, html_content

    def delete_pic(self):
        try:
            _pic_path = os.path.join(os.getcwd(), "doc", "pic")
            if not os.path.exists(_pic_path):
                os.makedirs(_pic_path)
            for item in os.listdir(_pic_path):
                os.remove(os.path.join(_pic_path, item))
        except Exception as e:
            print("delete_pic", e)

    def get_empty_params(self):
        return collections.OrderedDict()

    def report_basic(self, results):
        self.base_info = results

    def report_sn(self, sn):
        if self.sn == "":
            self.sn = sn
            self.write_content(f"设备序号：{self.sn}")

    def report_start(self, start_time):
        if self.start_time == 0:
            path = os.path.join(os.getcwd()) + "\\libutils\\default.docx"
            print(path)
            self.document = Document(path)
            self.write_title("基本信息", 1)
            if self.base_info is not None:
                self.write_report_map(self.base_info)
            self.start_time = start_time
            file_create_time = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(start_time))
            self.write_content(f"测试时间：{file_create_time}")

    def report_video(self, results):
        self.write_title("图像云台测试", 1)
        self.write_report_map(results)

    def report_extern(self, results):
        self.write_title("外部接口测试", 1)
        self.write_report_map(results)

    def report_end(self, result, end):
        self.write_title("测试结果", 1)
        self.write_content(f"测试结果：{result}, 测试耗时：{int(end - self.start_time)}s")

    def write_report_map(self, results):
        for item_key in results.keys():
            item_value = results.get(item_key, "").strip()
            if item_value.endswith((".jpg", ".png")):
                self.write_content(f"{item_key}：")
                self.write_pic(item_value)
            else:
                self.write_content(f"{item_key}：{item_value}")
            print("report:", item_key, results.get(item_key))

    def write_title(self, title, level):
        self.document.add_heading(title, level)

    def write_content(self, content):
        self.document.add_paragraph(content)

    def write_pic(self, pic_path):
        try:
            pic = Image.open(pic_path)
            if pic.mode == "P":
                pic = pic.convert("RGB")
            pic.save(pic_path, quality=100)
            self.document.add_picture(pic_path, width=Cm(15.2))
        except:
            self.write_content(f"图片获取失败：{pic_path}")

    def write_list(self, header_list, data, style="Table Grid"):
        # row行, col列
        table_col = len(header_list)
        # 新建一个1行n列的表头
        table = self.document.add_table(1, table_col)
        header_cells = table.rows[0].cells
        for i in range(len(header_cells)):
            header_cells[i].text = header_list[i]

        for item in data:
            row_cells = table.add_row().cells
            for i in range(len(row_cells)):
                row_cells[i].text = str(item[i])

        table.style = style

    def write_page_break(self):
        self.document.add_page_break()


report_doc = WordReport()

# [Python-docx使用](https://zhuanlan.zhihu.com/p/82880510)
# [python-docx官方文档](https://python-docx.readthedocs.io/en/latest/)
