#!/usr/bin/env python3
# -*- coding: utf-8 -*-

# year = 8200 * 12
# total = 0
# for i in range(1, 31):
#     total += year
#     print(f"第{i}年, year：{year}, total:{total}")
#     year *= 0.977
import base64
import json

import docx

# doc = docx.Document("tcp_interface.docx")
# last_line = ""
#
# cache = ""
# dep = -1
# index = 1
# 数据提取
# for p in doc.paragraphs:
#     graph = p.text
#     # print(graph)
#     if len(graph) > 0:
#         if graph.startswith("请求命令参数"):
#             dep = 0
#             print(last_line)
#             # print(graph)
#             continue
#         if dep >= 0:
#             cache += graph
#             json_start = graph.count("{")
#             json_end = graph.count("}")
#             dep = dep + json_start - json_end
#             # print(dep, graph)
#             if dep == 0:
#                 cache = cache.replace("”", "\"").replace(" ", "").replace("\t", "")
#                 print(index, cache)
#                 dep = -1
#                 cache = ""
#                 index += 1
#             continue
#         last_line = graph

# 暂存文件
# with open("tcp_interface_auto_cmd.txt", "w") as of:
#     for p in doc.paragraphs:
#         graph = p.text
#         # print(graph)
#         if len(graph) > 0:
#             if graph.startswith("请求命令参数"):
#                 dep = 0
#                 print(last_line)
#                 # print(graph)
#                 continue
#             if dep >= 0:
#                 cache += graph
#                 json_start = graph.count("{")
#                 json_end = graph.count("}")
#                 dep = dep + json_start - json_end
#                 # print(dep, graph)
#                 if dep == 0:
#                     cache = cache.replace("”", "\"").replace(" ", "").replace("\t", "")
#                     of.write(cache)
#                     print(index, cache)
#                     of.write('\n')
#                     dep = -1
#                     cache = ""
#                     index += 1
#                 continue
#             last_line = graph

# cmd = '{"body":{"recognition_area":{"polygon_num":1,"polygon":[{"id":1,"enable":true,"point_num":4,"point":[{"x":4096,"y":4096},{"x":12288,"y":4096},{"x":13926,"y":13926},{"x":2457,"y":13926}]}]}},"cmd":"set_reco_para","id":"132156"}'
# # cmd = '{"cmd":"get_alg_para","body":{"rule_chn":0,"alg_prm_type":"car_motion_prm"},"id":"132158"}'
# # cmd = '{"cmd":"playserver_json_request","id":"132156","body":{"type":"ps_get_voice_info","voice_type":3}}'
# # cmd = '{"cmd":"get_traffic_lights"}'
#
# params = []
# func_note = ""
# func_name = ""
# func_body = ""
# func_end = ""
#
#
# def json_parser(root, json_obj):
#     global func_name, params, func_end
#     try:
#         if isinstance(json_obj, dict):
#             print("dict:", len(json_obj), root, json_obj)
#             for key, value in json_obj.items():
#                 if key == "id":
#                     continue
#                 elif key == "cmd":
#                     func_name = value
#                     if func_name.startswith("get"):
#                         func_end = "return data"
#                 else:
#                     json_parser(key, value)
#         elif isinstance(json_obj, list):
#             print("list:", len(json_obj), root, json_obj)
#             params.append(root)
#             # for item in json_obj:
#             #     json_parser(root, item)
#         else:
#             params.append(root)
#             print("value:", root, json_obj)
#     except Exception as e:
#         print(e)
#
#
# # def create_func_body(json_obj, params_b):
# #     try:
# #         body_params = json_obj["body"]
# #
# #
# #     except Exception as e:
# #         print(e)
#
# json_parser("", json.loads(cmd))
# # func_body = create_func_body(json.loads(cmd), params)
# print(func_name, params, func_body, func_end)
#
# print(
# f'''
# def {func_name}(self, cmd_string_data=None):
#     """
#     {func_note}
#     :return:
#     """
#     if cmd_string_data is None:
#         cmd_string_data = '{cmd}'
#     data = self.client.tcp_send_cmd(cmd_string_data)[0]
#     assert data.get('cmd') == '{func_name}'
#     assert data.get('state_code') == 200
#     {func_end}
# '''
# )
#
# # print(
# # f"""
# # def {func_name}(self, {", ".join(params)}):
# #      cmd_string = "{func_name}"
# #      {func_body}
# #      data = self.client.tcp_send_cmd(cmd_string_data)[0]
# #      assert data.get('cmd') == cmd_string
# #      assert data.get('state_code') == 200
# #      {func_end}
# # """
# # )


# for i in range(5):
#     print(i)
# else:
#     print("for -> else")


a = base64.b64encode("左".encode('utf-8'))
print(a.decode('utf-8'))
